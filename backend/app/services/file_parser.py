import os
import platform
import re
import subprocess
import sys
import tempfile

from docx import Document


class FileParser:
    @staticmethod
    def parse_md(content: str) -> str:
        return content

    @staticmethod
    def parse_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
        except Exception as e:
            raise ValueError(f"无法打开docx文件: {str(e)}")

        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style.name.startswith("Heading"):
                    level = para.style.name
                    if "1" in level:
                        paragraphs.append(f"# {text}")
                    elif "2" in level:
                        paragraphs.append(f"## {text}")
                    elif "3" in level:
                        paragraphs.append(f"### {text}")
                    else:
                        paragraphs.append(f"#### {text}")
                else:
                    paragraphs.append(text)

        for table in doc.tables:
            paragraphs.append("\n| 列1 | 列2 | 列3 |")
            paragraphs.append("|------|------|------|")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(f"| {' | '.join(cells)} |")

        return "\n\n".join(paragraphs)

    @staticmethod
    def parse_doc_with_word_com(file_path: str) -> str:
        """
        使用 Word COM 接口解析 .doc 文件 (Windows only)
        需要安装 Microsoft Word
        """
        if platform.system() != "Windows":
            return None

        try:
            import win32com.client
        except ImportError:
            return None

        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(abs_path, ReadOnly=True)

            text = doc.Content.Text.strip()
            return text
        except Exception:
            return None
        finally:
            if doc:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass

    @staticmethod
    def parse_doc_convert_to_docx(file_path: str) -> str:
        """
        将 .doc 文件转换为 .docx 后解析 (Windows only)
        需要安装 Microsoft Word
        """
        if platform.system() != "Windows":
            return None

        try:
            import win32com.client
        except ImportError:
            return None

        word = None
        doc = None
        temp_file = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(abs_path, ReadOnly=True)

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                temp_file = f.name

            wdFormatXMLDocument = 16
            doc.SaveAs2(temp_file, FileFormat=wdFormatXMLDocument)
            doc.Close(False)
            doc = None

            result = FileParser.parse_docx(temp_file)
            return result
        except Exception:
            return None
        finally:
            if doc:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            if temp_file and os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except OSError:
                    pass

    @staticmethod
    def parse_doc_with_olefile(file_path: str) -> str:
        """
        使用 olefile 库解析 .doc 文件
        这个方法只能提取纯文本，会丢失格式
        """
        try:
            import olefile
        except ImportError:
            return None

        try:
            ole = olefile.OleFileIO(file_path)
            if ole.exists("WordDocument"):
                stream = ole.openstream("WordDocument")
                data = stream.read()

                text_parts = []
                for i in range(len(data)):
                    if 32 <= data[i] < 127:
                        text_parts.append(chr(data[i]))
                    elif data[i] in (10, 13):
                        text_parts.append("\n")

                text = "".join(text_parts)
                text = re.sub(r"\n{3,}", "\n\n", text)
                return text.strip() if text.strip() else None
        except Exception:
            pass
        return None

    @staticmethod
    def parse_doc(file_path: str) -> str:
        """
        解析.doc文件
        尝试使用多种方法解析.doc文件
        """
        errors = []

        if platform.system() == "Windows":
            result = FileParser.parse_doc_convert_to_docx(file_path)
            if result:
                return result
            errors.append("Word COM 转换方式失败")

            result = FileParser.parse_doc_with_word_com(file_path)
            if result:
                return result
            errors.append("Word COM 直接读取方式失败")

        try:
            result = subprocess.run(
                ["antiword", file_path], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            errors.append(f"antiword: {str(e)}")
        except Exception as e:
            errors.append(f"antiword: {str(e)}")

        try:
            result = subprocess.run(
                ["catdoc", file_path], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            errors.append(f"catdoc: {str(e)}")
        except Exception as e:
            errors.append(f"catdoc: {str(e)}")

        result = FileParser.parse_doc_with_olefile(file_path)
        if result:
            return result
        errors.append("olefile 方式失败")

        try:
            return FileParser.parse_docx(file_path)
        except Exception as e:
            errors.append(f"python-docx: {str(e)}")

        try:
            import textract

            text = textract.process(file_path).decode("utf-8")
            return text.strip()
        except ImportError:
            errors.append("textract 未安装")
        except Exception as e:
            errors.append(f"textract: {str(e)}")

        error_detail = "; ".join(errors)
        return f"[无法解析.doc文件内容，请将文件转换为.docx格式或手动复制内容]\n\n文件路径: {file_path}\n解析尝试: {error_detail}"

    @staticmethod
    def parse_file(file_path: str, file_extension: str) -> str:
        ext = file_extension.lower().replace(".", "")

        if ext == "md" or ext == "markdown":
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        elif ext == "docx":
            return FileParser.parse_docx(file_path)
        elif ext == "doc":
            return FileParser.parse_doc(file_path)
        elif ext == "txt":
            with open(file_path, encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
